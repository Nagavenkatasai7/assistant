"""
DOCX generation engine - converts markdown resume to ATS-friendly Word document
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import re


class DOCXGenerator:
    """Generate ATS-friendly DOCX from markdown resume"""

    def __init__(self):
        """Initialize DOCX generator"""
        pass

    def markdown_to_docx(self, markdown_content, output_path, company_name=None):
        """
        Convert markdown resume to ATS-friendly DOCX format

        Args:
            markdown_content: Resume in markdown format
            output_path: Path to save DOCX file
            company_name: Optional company name for filename

        Returns:
            Path to generated DOCX file
        """
        # Create document
        doc = Document()

        # Set margins (0.75 inches on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Setup custom styles
        self._setup_styles(doc)

        # Parse markdown and add to document
        self._parse_markdown_to_docx(doc, markdown_content)

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return str(output_path)

    def _setup_styles(self, doc):
        """Setup custom ATS-friendly styles"""
        styles = doc.styles

        # Name style (Heading 1 customization)
        if 'Resume Name' not in styles:
            name_style = styles.add_style('Resume Name', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.name = 'Calibri'
            name_style.font.size = Pt(16)
            name_style.font.bold = True
            name_style.font.color.rgb = RGBColor(0, 0, 0)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_style.paragraph_format.space_after = Pt(6)

        # Contact info style
        if 'Resume Contact' not in styles:
            contact_style = styles.add_style('Resume Contact', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.name = 'Calibri'
            contact_style.font.size = Pt(10)
            contact_style.font.color.rgb = RGBColor(0, 0, 0)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(12)

        # Section header style
        if 'Resume Section Header' not in styles:
            section_style = styles.add_style('Resume Section Header', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(12)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0, 0, 0)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)

        # Job title style
        if 'Resume Job Title' not in styles:
            job_style = styles.add_style('Resume Job Title', WD_STYLE_TYPE.PARAGRAPH)
            job_style.font.name = 'Calibri'
            job_style.font.size = Pt(11)
            job_style.font.bold = True
            job_style.font.color.rgb = RGBColor(0, 0, 0)
            job_style.paragraph_format.space_after = Pt(3)

        # Normal text style
        if 'Resume Normal' not in styles:
            normal_style = styles.add_style('Resume Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(10)
            normal_style.font.color.rgb = RGBColor(0, 0, 0)
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.15

    def _parse_markdown_to_docx(self, doc, markdown_content):
        """Parse markdown content and add to DOCX document"""
        lines = markdown_content.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # H1 - Candidate Name
            if line.startswith('# '):
                name = line[2:].strip()
                p = doc.add_paragraph(name, style='Resume Name')
                i += 1
                continue

            # H2 - Section Headers
            if line.startswith('## '):
                header = line[3:].strip()
                p = doc.add_paragraph(header.upper(), style='Resume Section Header')
                i += 1
                continue

            # H3 - Job Titles/Subsections
            if line.startswith('### '):
                title = line[4:].strip()
                p = doc.add_paragraph(title, style='Resume Job Title')
                i += 1
                continue

            # Italic text (usually dates/locations)
            if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                text = line[1:-1].strip()
                p = doc.add_paragraph(style='Resume Normal')
                run = p.add_run(text)
                run.italic = True
                i += 1
                continue

            # Bullet points
            if line.startswith('- ') or line.startswith('• '):
                bullet_text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                self._add_formatted_text(p, bullet_text)
                i += 1
                continue

            # Pipe separator (contact info line)
            if '|' in line and i < 5:  # Contact info usually in first few lines
                contact_line = line.replace('|', ' | ')
                p = doc.add_paragraph(style='Resume Contact')
                self._add_formatted_text(p, contact_line, add_hyperlinks=True)
                i += 1
                continue

            # Regular paragraph
            if line:
                p = doc.add_paragraph(style='Resume Normal')
                self._add_formatted_text(p, line)

            i += 1

    def _add_formatted_text(self, paragraph, text, add_hyperlinks=False):
        """
        Add text with markdown formatting to paragraph

        Handles:
        - Bold: **text**
        - Italic: *text*
        - Links: [text](url) (if add_hyperlinks=True)
        """
        # Handle hyperlinks in contact info
        if add_hyperlinks:
            # Simple hyperlink patterns for common sites
            hyperlink_patterns = [
                (r'LinkedIn', 'https://linkedin.com/in/naga-venkata-sai-chennu'),
                (r'GitHub', 'https://github.com/nagavenkatasai7'),
                (r'Portfolio', 'https://nagavenkatasai7.github.io/portfolio/'),
            ]

            for pattern, url in hyperlink_patterns:
                if pattern in text:
                    parts = text.split(pattern)
                    for idx, part in enumerate(parts[:-1]):
                        paragraph.add_run(part)
                        self._add_hyperlink(paragraph, pattern, url)
                    paragraph.add_run(parts[-1])
                    return

        # Process bold and italic formatting
        current_pos = 0
        text_length = len(text)

        while current_pos < text_length:
            # Check for bold (**text**)
            bold_match = re.search(r'\*\*([^*]+)\*\*', text[current_pos:])
            # Check for italic (*text*)
            italic_match = re.search(r'\*([^*]+)\*', text[current_pos:])

            # Find the earliest match
            next_match = None
            match_type = None

            if bold_match and italic_match:
                if bold_match.start() < italic_match.start():
                    next_match = bold_match
                    match_type = 'bold'
                else:
                    next_match = italic_match
                    match_type = 'italic'
            elif bold_match:
                next_match = bold_match
                match_type = 'bold'
            elif italic_match:
                next_match = italic_match
                match_type = 'italic'

            if next_match:
                # Add text before match
                before_text = text[current_pos:current_pos + next_match.start()]
                if before_text:
                    paragraph.add_run(before_text)

                # Add formatted text
                formatted_text = next_match.group(1)
                run = paragraph.add_run(formatted_text)
                if match_type == 'bold':
                    run.bold = True
                elif match_type == 'italic':
                    run.italic = True

                current_pos += next_match.end()
            else:
                # No more matches, add remaining text
                paragraph.add_run(text[current_pos:])
                break

    def _add_hyperlink(self, paragraph, text, url):
        """
        Add hyperlink to paragraph

        Simplified approach: just add styled text without actual hyperlink
        since python-docx hyperlink support is complex
        """
        # For now, just add the text with underline styling
        # Full hyperlink implementation requires complex XML manipulation
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        run.underline = True


def main():
    """Test DOCX generator"""
    sample_markdown = """# John Doe
johndoe@email.com | (555) 123-4567 | LinkedIn | GitHub | Portfolio | San Francisco, CA

## PROFESSIONAL SUMMARY
Senior Software Engineer with 5+ years of experience in AI/ML and full-stack development. Specialized in building scalable RAG systems and LLM applications using Python, LangChain, and cloud technologies.

## TECHNICAL SKILLS
**Programming:** Python, JavaScript, TypeScript, SQL
**AI/ML:** LangChain, OpenAI, Anthropic Claude, RAG, Prompt Engineering
**Cloud:** AWS, Docker, Kubernetes
**Databases:** PostgreSQL, MongoDB, Pinecone

## PROFESSIONAL EXPERIENCE

### Senior AI Engineer | TechCorp | San Francisco, CA
*01/2022 - Present*
- Built production RAG system handling 10K+ queries/day with 95% accuracy
- Implemented LLM-powered features using GPT-4 and Claude, reducing support costs by 40%
- Led team of 3 engineers in developing AI platform serving 50K+ users

### Software Engineer | StartupXYZ | Remote
*06/2019 - 12/2021*
- Developed full-stack applications using React and Node.js
- Implemented CI/CD pipelines reducing deployment time by 60%

## EDUCATION
**Master of Science in Computer Science** | Stanford University | 2019
**Bachelor of Science in Computer Science** | UC Berkeley | 2017

## CERTIFICATIONS
AWS Certified Solutions Architect - Professional
"""

    generator = DOCXGenerator()
    output_path = "test_resume.docx"

    try:
        result_path = generator.markdown_to_docx(sample_markdown, output_path)
        print(f"✓ DOCX generated successfully: {result_path}")
    except Exception as e:
        print(f"✗ DOCX generation failed: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
