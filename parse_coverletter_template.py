"""
Parse cover letter template DOCX for formatting guide
"""
from docx import Document
from pathlib import Path
import json

def parse_template(template_path):
    """Extract structure and formatting from DOCX template"""
    try:
        doc = Document(template_path)

        template_info = {
            "filename": Path(template_path).name,
            "paragraphs": [],
            "structure": {
                "sections": [],
                "formatting_notes": []
            }
        }

        print(f"Parsing template: {Path(template_path).name}")
        print(f"Total paragraphs: {len(doc.paragraphs)}\n")

        current_section = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if text:
                # Detect section headers (typically bold or larger font)
                is_header = False
                if para.runs:
                    is_header = para.runs[0].bold or (para.runs[0].font.size and para.runs[0].font.size.pt > 12)

                para_info = {
                    "index": i,
                    "text": text[:100] + "..." if len(text) > 100 else text,
                    "is_header": is_header,
                    "alignment": str(para.alignment) if para.alignment else "LEFT"
                }

                template_info["paragraphs"].append(para_info)

                if is_header:
                    current_section = text
                    template_info["structure"]["sections"].append({
                        "name": text,
                        "index": i
                    })
                    print(f"📍 Section {len(template_info['structure']['sections'])}: {text}")

        # Extract key formatting patterns
        template_info["structure"]["formatting_notes"] = [
            "Standard business letter format",
            "Left-aligned text",
            "Single spacing within paragraphs",
            "Double spacing between paragraphs",
            "Professional font (Arial, Calibri, or Times New Roman)",
            "11-12pt font size",
            "1-inch margins"
        ]

        return template_info

    except Exception as e:
        print(f"Error parsing template: {e}")
        return None

def main():
    template_path = "coverletter-knowledge/Professional-cover-letter-template2-1.docx"

    if not Path(template_path).exists():
        print(f"❌ Template not found: {template_path}")
        return

    template_info = parse_template(template_path)

    if template_info:
        # Save parsed template info
        output_file = "coverletter_template_structure.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(template_info, f, indent=2)

        print(f"\n✓ Template structure saved to {output_file}")
        print(f"✓ Found {len(template_info['structure']['sections'])} sections")
        print(f"✓ Total paragraphs: {len(template_info['paragraphs'])}")
    else:
        print("❌ Failed to parse template")

if __name__ == "__main__":
    main()
